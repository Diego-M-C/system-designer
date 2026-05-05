#!/usr/bin/env python3
"""Generate professionally formatted .docx reports for all 16 Anthropic Academy courses."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from datetime import date

OUTPUT_DIR = "/mnt/c/Users/Diego/Desktop/AGENC_IA/Informes_Cursos_Anthropic"

# Brand colors
BRAND_ORANGE = RGBColor(0xD9, 0x73, 0x06)
BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_GRAY = RGBColor(0x55, 0x55, 0x55)
BRAND_LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)
BRAND_BODY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val["sz"]}" '
                f'w:space="0" w:color="{val["color"]}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(8)
    for run in h.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(22)
            run.font.color.rgb = BRAND_DARK
        elif level == 2:
            run.font.size = Pt(16)
            run.font.color.rgb = BRAND_DARK
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = BRAND_ORANGE
    return h


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("\u2500" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)


def add_body_text(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_BODY
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = BRAND_BODY
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_BODY
    return p


def add_info_table(doc, rows_data):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row_idx, (label, value) in enumerate(rows_data):
        row = table.rows[row_idx]
        cell_l = row.cells[0]
        cell_v = row.cells[1]

        # Set widths
        cell_l.width = Cm(5)
        cell_v.width = Cm(12)

        # Label cell
        cell_l.paragraphs[0].clear()
        run = cell_l.paragraphs[0].add_run(label)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_DARK
        cell_l.paragraphs[0].paragraph_format.space_before = Pt(4)
        cell_l.paragraphs[0].paragraph_format.space_after = Pt(4)
        cell_l.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell_l, 'F0EBE3')

        # Value cell
        cell_v.paragraphs[0].clear()
        run = cell_v.paragraphs[0].add_run(value)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_BODY
        cell_v.paragraphs[0].paragraph_format.space_before = Pt(4)
        cell_v.paragraphs[0].paragraph_format.space_after = Pt(4)
        cell_v.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Borders
        border_spec = {"val": "single", "sz": "4", "color": "E0D8CE"}
        for cell in [cell_l, cell_v]:
            set_cell_borders(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)

    return table


def add_header_footer(doc, title):
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(f"Anthropic Academy  |  {title}")
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = BRAND_LIGHT_GRAY
        run.font.italic = True

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("Professional Course Report  |  Source: anthropic.skilljar.com  |  ")
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = BRAND_LIGHT_GRAY
        run = fp.add_run(date.today().strftime("%B %Y"))
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.font.color.rgb = BRAND_LIGHT_GRAY


def create_report(course_data):
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Default style --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BRAND_BODY
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.25

    # ========== COVER PAGE ==========
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Orange accent line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2588" * 40)
    run.font.size = Pt(4)
    run.font.color.rgb = BRAND_ORANGE

    doc.add_paragraph("")

    # Academy branding
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("A N T H R O P I C    A C A D E M Y")
    run.font.size = Pt(12)
    run.font.color.rgb = BRAND_ORANGE
    run.font.name = 'Calibri'
    run.bold = True
    # letter spacing via XML
    rPr = run._element.get_or_add_rPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:val="40"/>')
    rPr.append(spacing)

    doc.add_paragraph("")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(course_data['title'].upper())
    run.font.size = Pt(26)
    run.font.color.rgb = BRAND_DARK
    run.font.name = 'Calibri'
    run.bold = True

    # Subtitle line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Professional Course Report")
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_GRAY
    run.font.name = 'Calibri'

    doc.add_paragraph("")

    # Orange accent line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2588" * 40)
    run.font.size = Pt(4)
    run.font.color.rgb = BRAND_ORANGE

    for _ in range(3):
        doc.add_paragraph("")

    # Meta information block
    meta_items = [
        ("Platform:", "Anthropic Academy (Skilljar)"),
        ("Source:", course_data.get('url', 'https://anthropic.skilljar.com/')),
        ("Report Date:", date.today().strftime("%B %d, %Y")),
        ("Classification:", "Public Course Information"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label + "  ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BRAND_GRAY
        run.font.name = 'Calibri'
        run = p.add_run(value)
        run.font.size = Pt(9)
        run.font.color.rgb = BRAND_GRAY
        run.font.name = 'Calibri'

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    add_styled_heading(doc, "Table of Contents", level=1)
    add_separator(doc)

    toc_items = [
        "1.  Executive Summary",
        "2.  Course Overview",
        "3.  Learning Objectives",
        "4.  Prerequisites",
        "5.  Target Audience",
        "6.  Course Content & Curriculum",
        "7.  Technical Metadata",
        "8.  Disclaimer & Methodology",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(item)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = BRAND_DARK

    doc.add_page_break()

    # Add header/footer
    add_header_footer(doc, course_data['title'])

    # ========== 1. EXECUTIVE SUMMARY ==========
    add_styled_heading(doc, "1. Executive Summary", level=1)
    add_separator(doc)
    add_body_text(doc, course_data.get('executive_summary', 'N/A'))

    # ========== 2. COURSE OVERVIEW ==========
    add_styled_heading(doc, "2. Course Overview", level=1)
    add_separator(doc)

    doc.add_paragraph("")
    overview_fields = [
        ("Course Title", course_data['title']),
        ("Short Description", course_data.get('short_description', 'N/A')),
        ("Course ID", course_data.get('course_id', 'N/A')),
        ("Cost", course_data.get('cost', 'Free')),
        ("Format", "Online, Self-Paced"),
        ("Platform", "Anthropic Academy (Skilljar LMS)"),
        ("Source URL", course_data.get('url', 'N/A')),
    ]
    add_info_table(doc, overview_fields)

    if course_data.get('long_description'):
        doc.add_paragraph("")
        add_body_text(doc, course_data['long_description'])

    # ========== 3. LEARNING OBJECTIVES ==========
    add_styled_heading(doc, "3. Learning Objectives", level=1)
    add_separator(doc)

    objectives = course_data.get('learning_objectives', [])
    if objectives:
        add_body_text(doc, "Upon completion of this course, participants will be able to:")
        doc.add_paragraph("")
        for i, obj in enumerate(objectives, 1):
            if isinstance(obj, dict):
                title = obj.get('title', '')
                desc = obj.get('description', '')
                add_bullet(doc, desc, bold_prefix=f"{title}:")
            else:
                add_bullet(doc, obj)
    else:
        add_body_text(doc, "Detailed learning objectives are not publicly listed on the Anthropic Academy platform for this course. Specific objectives are available upon enrollment.", italic=True)

    # ========== 4. PREREQUISITES ==========
    add_styled_heading(doc, "4. Prerequisites", level=1)
    add_separator(doc)

    prereqs = course_data.get('prerequisites', [])
    if prereqs:
        for prereq in prereqs:
            add_bullet(doc, prereq)
    else:
        add_body_text(doc, "No specific prerequisites are publicly listed for this course on the Anthropic Academy platform.")

    # ========== 5. TARGET AUDIENCE ==========
    add_styled_heading(doc, "5. Target Audience", level=1)
    add_separator(doc)

    audience = course_data.get('target_audience', [])
    if audience:
        add_body_text(doc, "This course is designed for:")
        for a in audience:
            add_bullet(doc, a)
    else:
        add_body_text(doc, "Target audience details are not explicitly listed on the course page.")

    # ========== 6. COURSE CONTENT & CURRICULUM ==========
    add_styled_heading(doc, "6. Course Content & Curriculum", level=1)
    add_separator(doc)

    content = course_data.get('curriculum_details', '')
    if content:
        add_body_text(doc, content)

    # Topics
    topics = course_data.get('topics', [])
    if topics:
        doc.add_paragraph("")
        add_styled_heading(doc, "Key Topics Covered", level=3)
        for topic in topics:
            add_bullet(doc, topic)

    # Resources
    resources = course_data.get('resources', [])
    if resources:
        doc.add_paragraph("")
        add_styled_heading(doc, "Supplementary Resources", level=3)
        for r in resources:
            add_bullet(doc, r)

    # Certification
    if course_data.get('certification'):
        doc.add_paragraph("")
        add_styled_heading(doc, "Certification", level=3)
        add_body_text(doc, course_data['certification'])

    # ========== 7. TECHNICAL METADATA ==========
    add_styled_heading(doc, "7. Technical Metadata", level=1)
    add_separator(doc)

    doc.add_paragraph("")
    tech_fields = [
        ("Course ID", course_data.get('course_id', 'N/A')),
        ("Published Course ID", course_data.get('published_course_id', 'N/A')),
        ("Platform", "Skilljar Learning Management System"),
        ("Timezone", "America/Los_Angeles"),
        ("Source URL", course_data.get('url', 'N/A')),
        ("Access", course_data.get('cost', 'Free')),
    ]
    add_info_table(doc, tech_fields)

    # ========== 8. DISCLAIMER ==========
    doc.add_page_break()
    add_styled_heading(doc, "8. Disclaimer & Methodology", level=1)
    add_separator(doc)

    disclaimer_text = (
        "This report was generated based exclusively on publicly available information "
        "from the Anthropic Academy platform (https://anthropic.skilljar.com/). "
        "No content has been fabricated or invented. All descriptions, learning objectives, "
        "prerequisites, and target audience information are extracted directly from the "
        "official course pages as published by Anthropic."
    )
    add_body_text(doc, disclaimer_text, italic=True)

    methodology_text = (
        "Where specific details such as learning objectives, prerequisites, or curriculum "
        "breakdowns are not publicly available on the course landing page, this has been "
        "explicitly noted. Some courses may contain additional content visible only after "
        "enrollment that is not reflected in this report."
    )
    add_body_text(doc, methodology_text, italic=True)

    accuracy_text = (
        f"Information is accurate as of the report generation date: {date.today().strftime('%B %d, %Y')}. "
        "Course content may be updated by Anthropic at any time."
    )
    add_body_text(doc, accuracy_text, italic=True)

    # Save
    filename = course_data['filename']
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"  [OK] {filename}")
    return filepath


# ===== ALL 16 COURSES DATA =====
courses = [
    {
        "title": "Building with the Claude API",
        "filename": "01_Building_with_the_Claude_API.docx",
        "url": "https://anthropic.skilljar.com/claude-with-the-anthropic-api/",
        "short_description": "This comprehensive course covers the full spectrum of working with Anthropic models using the Claude API.",
        "long_description": "This comprehensive video course teaches developers how to integrate Claude AI into applications using the Anthropic API. It covers API setup and authentication, conversation implementation, prompt engineering, tool integration, RAG systems, extended thinking, multimodal features, prompt caching, MCP servers, and agent-based systems.",
        "course_id": "ca6k4fml7abo",
        "published_course_id": "338y0cdn1ivc3",
        "cost": "Free",
        "executive_summary": "\"Building with the Claude API\" is a comprehensive technical course offered by the Anthropic Academy that provides end-to-end training on integrating Claude AI into applications via the Anthropic API. The course targets developers across multiple disciplines and covers the full development lifecycle from API setup to production-ready agent architectures. It represents Anthropic's primary developer education resource for API-level integration.",
        "learning_objectives": [
            {"title": "API Setup and Authentication", "description": "Including key management and initial configuration"},
            {"title": "Single and Multi-turn Conversation Implementation", "description": "Building conversational flows with proper context handling"},
            {"title": "System Prompts and Model Behavior Configuration", "description": "Configuring Claude's behavior through system-level instructions"},
            {"title": "Prompt Evaluation Workflows", "description": "Using test datasets to evaluate and iterate on prompts"},
            {"title": "Prompt Engineering Techniques", "description": "XML structuring, example-based learning, and advanced prompting strategies"},
            {"title": "Tool Use Integration", "description": "Custom tool development and function calling capabilities"},
            {"title": "Retrieval-Augmented Generation (RAG) Systems", "description": "Building knowledge-grounded applications"},
            {"title": "Extended Thinking Mode and Multimodal Features", "description": "Leveraging Claude's reasoning and vision capabilities"},
            {"title": "Prompt Caching Strategies", "description": "Optimizing performance and cost through caching"},
            {"title": "Model Context Protocol (MCP) Servers", "description": "Building and integrating MCP servers for extended functionality"},
            {"title": "Anthropic Apps Deployment", "description": "Including Claude Code for automated development tasks"},
            {"title": "Agent-Based Systems", "description": "Parallelization, routing, and orchestration patterns for complex AI workflows"},
        ],
        "prerequisites": [
            "Python programming proficiency",
            "Basic JSON handling knowledge",
        ],
        "target_audience": [
            "Backend developers building AI-powered APIs",
            "Full-stack engineers integrating conversational AI",
            "Data engineers implementing document processing",
            "DevOps professionals automating workflows",
            "Technical architects designing AI systems",
            "Software engineers transitioning to AI/ML",
            "Chatbot and virtual assistant developers",
        ],
        "curriculum_details": "The course delivers video-based instruction covering the full spectrum of Claude API capabilities. Starting with foundational API setup and authentication, it progresses through conversation management, prompt engineering, tool use, RAG implementation, advanced features (extended thinking, multimodal, caching), MCP integration, and culminates with agent-based system architectures including parallelization and routing patterns.",
    },
    {
        "title": "Claude Code in Action",
        "filename": "02_Claude_Code_in_Action.docx",
        "url": "https://anthropic.skilljar.com/claude-code-in-action",
        "short_description": "Integrate Claude Code into your development workflow.",
        "long_description": "This course delivers comprehensive training on using Claude Code for software development, covering AI coding assistant architecture, practical implementation techniques, and advanced integration strategies. Students learn about context management approaches and how to extend functionality through MCP servers and GitHub integration.",
        "course_id": "3n2veylcj0hl",
        "published_course_id": "1ylsvrzjdwl86",
        "cost": "Free",
        "executive_summary": "\"Claude Code in Action\" is a hands-on course focused on integrating Claude Code into real-world software development workflows. It covers the architecture of AI coding assistants, tool use systems, context management, visual communication, custom automation, MCP server integration, and GitHub workflow automation. This course is essential for developers seeking to leverage Claude Code as a productivity multiplier in their daily engineering work.",
        "learning_objectives": [
            {"title": "Coding Assistant Architecture", "description": "Learn how AI assistants interact with codebases through tool integration and technical foundations enabling code analysis and modification"},
            {"title": "Claude Code's Tool Use System", "description": "Discover leveraging multiple tools in combination to handle complex, multi-step programming tasks across various development scenarios"},
            {"title": "Context Management Techniques", "description": "Learn strategies for maintaining relevant context throughout conversations and effectively referencing project resources for optimal AI assistance"},
            {"title": "Visual Communication Workflows", "description": "Understand using visual inputs to communicate interface changes and leverage advanced planning features for complex codebase modifications"},
            {"title": "Custom Automation", "description": "Explore building reusable custom commands and automations that streamline repetitive development tasks"},
            {"title": "MCP Server Integration", "description": "Learn to integrate external tools and services for enhanced capabilities like browser automation and specialized development workflows"},
            {"title": "GitHub Workflow Integration", "description": "Understand setting up automated code review processes and integrating AI assistance into existing version control workflows"},
            {"title": "Thinking and Planning Modes", "description": "Learn when and how to use different reasoning approaches for various complexity levels of programming challenges"},
        ],
        "prerequisites": [
            "Familiarity with command-line interfaces and terminal operations",
            "Basic understanding of version control with Git",
        ],
        "target_audience": [
            "Software developers looking to integrate AI assistance into their coding workflows",
            "Teams seeking to implement AI-powered GitHub integration for multiple workflows",
        ],
        "curriculum_details": "The course follows a progressive structure from understanding AI coding assistant architecture through practical tool use, context management, visual communication, custom automation creation, MCP server extension, and GitHub integration. It emphasizes hands-on application in real development scenarios.",
    },
    {
        "title": "Claude 101",
        "filename": "03_Claude_101.docx",
        "url": "https://anthropic.skilljar.com/claude-101",
        "short_description": "Learn how to use Claude for everyday work tasks, understand core features, and explore resources for more advanced learning on other topics.",
        "long_description": "Claude 101 is an introductory course designed to teach users how to use Claude for everyday work tasks. It covers core features of Claude and provides pathways to more advanced learning resources. The course uses video lessons, modular content, quizzes, and PDF resources.",
        "course_id": "22npsux5ldfq0",
        "published_course_id": "196eerrfd7ixv",
        "cost": "Free",
        "executive_summary": "\"Claude 101\" is the foundational entry point to the Anthropic Academy curriculum. It provides a practical introduction to using Claude for everyday work tasks, covering core features and capabilities. This course is ideal for users new to Claude who want a structured onboarding experience before progressing to more specialized courses.",
        "learning_objectives": [],
        "prerequisites": [],
        "target_audience": [
            "New Claude users seeking structured onboarding",
            "Professionals wanting to integrate Claude into everyday work tasks",
            "Anyone exploring AI assistants for productivity enhancement",
        ],
        "curriculum_details": "The course follows a modular structure using video lessons, quizzes, and PDF resources. It provides a progressive introduction to Claude's core features and everyday applications, with pathways to more advanced Anthropic Academy courses.",
        "topics": [
            "Core Claude features and capabilities",
            "Everyday work task applications",
            "Foundational AI interaction concepts",
            "Resources for advanced learning paths",
        ],
    },
    {
        "title": "Introduction to Claude Cowork",
        "filename": "04_Introduction_to_Claude_Cowork.docx",
        "url": "https://anthropic.skilljar.com/introduction-to-claude-cowork",
        "short_description": "Work alongside Claude on real files and projects.",
        "long_description": "This hands-on course teaches users to work alongside Claude on real files and projects. The platform focuses on the Cowork task loop, plugins and skills, file and research workflows, and responsible steering of multi-step work. As described in the course: \"Cowork is Claude working directly with your files, folders, and apps -- reading, editing, and producing real outputs on your machine.\"",
        "course_id": "okcui6s1t0",
        "published_course_id": "f47ytlftvz8z",
        "cost": "Free",
        "executive_summary": "\"Introduction to Claude Cowork\" provides hands-on training for working alongside Claude directly on files, folders, and applications. It covers the Cowork task loop, plugin and skill configuration, file and research workflows, and techniques for responsibly steering multi-step work. The course bridges the gap between conversational AI use and integrated, file-level collaboration.",
        "learning_objectives": [
            "Running end-to-end tasks with Claude Cowork",
            "Understanding how context shapes Claude's planning",
            "Setting up plugins and skills for specific workflows",
            "Steering longer-running tasks effectively",
        ],
        "prerequisites": [],
        "target_audience": [
            "Professionals seeking to integrate Claude into their daily workflows",
            "Users who want hands-on, file-level collaboration with Claude",
        ],
        "curriculum_details": "The course is designed to move participants from first launch to confident daily use. It covers the Cowork task loop, context management for planning, plugin and skill setup, and techniques for steering multi-step tasks. The emphasis is on practical, hands-on application rather than theoretical knowledge.",
    },
    {
        "title": "AI Fluency: Framework & Foundations",
        "filename": "05_AI_Fluency_Framework_and_Foundations.docx",
        "url": "https://anthropic.skilljar.com/ai-fluency-framework-foundations",
        "short_description": "Learn to collaborate with AI systems effectively, efficiently, ethically, and safely.",
        "long_description": "This course represents a partnership between Anthropic and academic experts Professor Joseph Feller (University College Cork) and Professor Rick Dakan (Ringling College). It teaches practical skills for effective, efficient, ethical, and safe AI interaction. The course has something for everyone, whether new to Claude or a seasoned AI practitioner.",
        "course_id": "17owe4fx9adox",
        "published_course_id": "3gyr99bvbs6cs",
        "cost": "Free",
        "executive_summary": "\"AI Fluency: Framework & Foundations\" is the cornerstone course of the AI Fluency series within the Anthropic Academy. Developed in partnership with academic experts Prof. Joseph Feller and Prof. Rick Dakan, it establishes the 4D AI Fluency Framework and teaches practical skills for effective, efficient, ethical, and safe AI interaction. This course serves as a prerequisite for several other AI Fluency courses in the academy.",
        "learning_objectives": [],
        "prerequisites": [],
        "target_audience": [
            "Anyone new to Claude or AI interaction",
            "Seasoned AI practitioners seeking structured frameworks",
            "Professionals in any field looking to develop AI fluency",
        ],
        "curriculum_details": "The course covers the 4D AI Fluency Framework, teaching practical skills for AI collaboration. It is developed in partnership with academic experts and includes downloadable PDF resources.",
        "resources": [
            "About this course (PDF)",
            "Meet your instructors (PDF)",
            "How we used AI in building the course (PDF)",
        ],
        "certification": "After finishing the course, students have the opportunity to take a final assessment and receive a certificate of completion.",
    },
    {
        "title": "Introduction to Model Context Protocol",
        "filename": "06_Introduction_to_Model_Context_Protocol.docx",
        "url": "https://anthropic.skilljar.com/introduction-to-model-context-protocol",
        "short_description": "Comprehensive coverage of the Model Context Protocol (MCP), focusing on building MCP servers and clients.",
        "long_description": "This course provides comprehensive coverage of the Model Context Protocol (MCP), focusing on building MCP servers and clients using the Python SDK. Students learn about MCP's three main primitives -- tools, resources, and prompts -- and how they integrate with Claude AI. The course covers MCP architecture, transport-agnostic communication, request-response flows, and practical integration patterns.",
        "course_id": "47ajyxsragmw",
        "published_course_id": "3ep4q1d5v7hue",
        "cost": "Free",
        "executive_summary": "\"Introduction to Model Context Protocol\" is a technical course that provides foundational training on MCP, Anthropic's open protocol for connecting AI systems to external tools and data sources. It covers MCP's three core primitives (tools, resources, prompts), server and client construction using the Python SDK, and practical integration patterns. This course is essential for developers building extensible AI applications.",
        "learning_objectives": [
            {"title": "MCP Architecture", "description": "Understanding how MCP transfers tool definition and execution to specialized servers"},
            {"title": "Transport-Agnostic Communication", "description": "Learning MCP's communication system and message types"},
            {"title": "Request-Response Flow", "description": "Understanding the complete flow from user queries to external services"},
            {"title": "Building MCP Servers", "description": "Constructing servers using the Python SDK with decorators"},
            {"title": "Document Management Functionality", "description": "Implementing tools for reading and editing operations"},
            {"title": "MCP Server Inspector", "description": "Using the built-in inspector for browser-based testing"},
            {"title": "Resource Definitions", "description": "Defining resources for read-only data, including static and templated URIs"},
            {"title": "Resource Reading in Clients", "description": "Implementing resource consumption with proper MIME type handling"},
            {"title": "Prompt Creation", "description": "Building pre-built, high-quality instructions as prompts"},
            {"title": "Primitive Differentiation", "description": "Knowing when to use each MCP primitive (tools vs. resources vs. prompts)"},
            {"title": "Autocomplete Integration", "description": "Practical integration patterns including autocomplete functionality"},
        ],
        "prerequisites": [
            "Working knowledge of Python programming",
            "Basic understanding of JSON and HTTP request-response patterns",
        ],
        "target_audience": [
            "Developers looking to build MCP servers",
            "Engineers integrating external tools and data with Claude AI",
        ],
        "curriculum_details": "The course progresses from MCP architecture fundamentals through the three core primitives (tools, resources, prompts), server construction with the Python SDK, client implementation, testing with the MCP Server Inspector, and practical integration patterns including autocomplete.",
    },
    {
        "title": "AI Fluency for Educators",
        "filename": "07_AI_Fluency_for_Educators.docx",
        "url": "https://anthropic.skilljar.com/ai-fluency-for-educators",
        "short_description": "Empowers faculty, instructional designers, and educational leaders to apply AI Fluency into their own teaching practice and institutional strategy.",
        "long_description": "This course empowers faculty, instructional designers, and educational leaders to apply AI Fluency into their own teaching practice and institutional strategy. Developed in partnership with Prof. Joseph Feller (University College Cork) and Prof. Rick Dakan (Ringling College).",
        "course_id": "f6as998uhezb",
        "published_course_id": "3hnpwx8d4az4y",
        "cost": "Free",
        "executive_summary": "\"AI Fluency for Educators\" extends the 4D AI Fluency Framework into educational practice and institutional strategy. Developed with academic experts, it targets faculty, instructional designers, and educational leaders who want to integrate AI fluency principles into their teaching. The course builds on the foundational AI Fluency course and includes certification upon completion.",
        "learning_objectives": [],
        "prerequisites": [
            "Foundational knowledge of the 4D AI Fluency Framework (covered in the core AI Fluency: Framework & Foundations course)",
        ],
        "target_audience": [
            "Academic faculty",
            "Instructional designers",
            "Educational leaders",
        ],
        "curriculum_details": "The course provides specialized training for applying the 4D AI Fluency Framework in educational contexts. It includes video lessons, modular content, quizzes, and PDF resources focused on teaching practice and institutional strategy.",
        "resources": [
            "About this course (PDF)",
            "Meet your instructors (PDF)",
            "How we used AI in building the course (PDF)",
        ],
        "certification": "After finishing the course, participants have the opportunity to take a final assessment and receive a certificate of completion.",
    },
    {
        "title": "AI Fluency for Students",
        "filename": "08_AI_Fluency_for_Students.docx",
        "url": "https://anthropic.skilljar.com/ai-fluency-for-students",
        "short_description": "Empowers students to develop AI Fluency skills that enhance learning, career planning, and academic success through responsible AI collaboration.",
        "long_description": "This course empowers students to develop AI Fluency skills that enhance learning, career planning, and academic success through responsible AI collaboration. It is a partnership between Anthropic and academic experts Prof. Joseph Feller (University College Cork) and Prof. Rick Dakan (Ringling College).",
        "course_id": "1yao8gc9rmfcz",
        "published_course_id": "3j3jd8xdnpp94",
        "cost": "Free",
        "executive_summary": "\"AI Fluency for Students\" applies the 4D AI Fluency Framework to the student experience, focusing on learning enhancement, career planning, and academic success through responsible AI collaboration. Developed with academic experts, it provides students with practical AI fluency skills within an ethical framework.",
        "learning_objectives": [],
        "prerequisites": [
            "Understanding of the 4D AI Fluency Framework covered in the core AI Fluency: Framework & Foundations course",
        ],
        "target_audience": [
            "Students seeking to enhance learning through responsible AI collaboration",
            "Students focused on career planning with AI skills",
            "Students aiming for academic success with AI tools",
        ],
        "curriculum_details": "The course adapts the 4D AI Fluency Framework for student use cases including learning enhancement, career planning, and academic success. It includes video lessons, modular content, quizzes, and PDF resources.",
        "resources": [
            "About this course (PDF)",
            "Meet your instructors (PDF)",
            "How we used AI in building the course (PDF)",
        ],
        "certification": "After finishing the course, students have the opportunity to take a final assessment and receive a certificate of completion.",
    },
    {
        "title": "Model Context Protocol: Advanced Topics",
        "filename": "09_MCP_Advanced_Topics.docx",
        "url": "https://anthropic.skilljar.com/model-context-protocol-advanced-topics",
        "short_description": "Discover advanced MCP implementation patterns including sampling, notifications, file system access, and transport mechanisms for production MCP server development.",
        "long_description": "This course examines advanced features and implementation patterns for MCP development, focusing on server-client communication, transport mechanisms, and production deployment considerations. It explores sophisticated capabilities including sampling for AI model integration, notification systems, file system access control, and technical details of different transport protocols.",
        "course_id": "322b89z3mttch",
        "published_course_id": "1ofs8qdfctxqp",
        "cost": "Free",
        "executive_summary": "\"Model Context Protocol: Advanced Topics\" is the advanced follow-up to the introductory MCP course, covering production-grade implementation patterns. It addresses sampling, notifications, file system access control, JSON message architecture, transport mechanisms (Stdio, StreamableHTTP), production scaling, and transport selection criteria. This course is critical for engineers building production MCP deployments.",
        "learning_objectives": [
            {"title": "Sampling Implementation", "description": "How MCP servers can request language model calls through connected clients, including architecture that shifts AI costs and complexity from server to client"},
            {"title": "Progress Notifications and Logging", "description": "Implementing real-time feedback systems using context objects, logging callbacks, and progress reporting for long-running operations"},
            {"title": "Roots-Based File Access", "description": "Permission systems granting MCP servers access to specific directories while providing security boundaries and friendly file discovery"},
            {"title": "JSON Message Architecture", "description": "Examining the full MCP message specification, distinguishing request-result pairs from notification messages, understanding bidirectional communication patterns"},
            {"title": "Stdio Transport Mechanisms", "description": "How MCP clients and servers communicate through standard input/output streams, including the required initialization handshake sequence"},
            {"title": "StreamableHTTP Transport", "description": "How Server-Sent Events (SSE) enable server-to-client communication over HTTP, including session management and dual-connection architectures"},
            {"title": "HTTP Transport Limitations", "description": "How configuration flags affect functionality regarding server-initiated requests and streaming capabilities"},
            {"title": "Production Scaling Considerations", "description": "When to use stateless HTTP for horizontal scaling with load balancers and trade-offs between stateful/stateless server configurations"},
            {"title": "Transport Selection Criteria", "description": "Choosing appropriate transport methods based on deployment requirements, functional needs, and scaling constraints"},
        ],
        "prerequisites": [
            "Experience with Python development and async programming patterns",
            "Familiarity with JSON message formats and HTTP protocols",
            "Basic knowledge of Server-Sent Events (SSE)",
        ],
        "target_audience": [
            "Developers working with Model Context Protocol implementations",
            "Engineers building MCP servers and clients for production environments",
        ],
        "curriculum_details": "The course progresses from advanced MCP capabilities (sampling, notifications, file access) through message architecture, transport mechanisms (Stdio, StreamableHTTP), transport limitations, production scaling patterns, and transport selection criteria. It focuses on production-ready implementation patterns.",
    },
    {
        "title": "Claude with Amazon Bedrock",
        "filename": "10_Claude_with_Amazon_Bedrock.docx",
        "url": "https://anthropic.skilljar.com/claude-in-amazon-bedrock",
        "short_description": "Comprehensive guide to integrating and deploying Claude AI models through Amazon Bedrock.",
        "long_description": "This technical course provides a comprehensive guide to integrating and deploying Claude AI models through Amazon Bedrock. Developers learn to implement Claude's API, build production-ready applications, and leverage advanced features including tool use, retrieval augmented generation (RAG), and autonomous agents. Originally created as part of an accreditation program for AWS, Anthropic launched this first-of-its-kind training for AWS employees and made it publicly available.",
        "course_id": "3i2kpf9wzkzfs",
        "published_course_id": "1rr2v7ii2ts1q",
        "cost": "Free",
        "executive_summary": "\"Claude with Amazon Bedrock\" is a technical course originally developed as an accreditation program for AWS employees, now publicly available through the Anthropic Academy. It provides comprehensive training on integrating and deploying Claude AI through Amazon Bedrock, covering API implementation, production application development, tool use, RAG, and autonomous agents within the AWS ecosystem.",
        "learning_objectives": [
            "Utilize Anthropic models on Amazon Bedrock for multi-turn conversations and system prompt configuration",
            "Build and evaluate prompts using structured approaches",
            "Design and integrate custom tools using JSON Schema for function calling and batch processing",
            "Develop RAG pipelines with text chunking, embeddings, BM25 search, and contextual retrieval techniques",
            "Configure and optimize Claude's advanced features including extended thinking, vision capabilities, and prompt caching",
            "Leverage Claude Code for automated debugging and task execution",
            "Implement Model Context Protocol (MCP) for defining tools, resources, and prompts in client applications",
            "Optimize inference through streaming, temperature control, and structured data extraction",
            "Build evaluation frameworks for prompts using model-based and code-based grading approaches",
        ],
        "prerequisites": [
            "Proficiency in Python programming",
            "Basic understanding of AWS services and Amazon Bedrock",
        ],
        "target_audience": [
            "Backend developers building AI-powered applications requiring advanced language model integration",
            "ML engineers implementing production RAG systems and conversational AI pipelines",
            "DevOps engineers deploying and optimizing Claude models in AWS infrastructure",
            "Full-stack developers creating applications with complex tool use and agent capabilities",
            "Technical architects designing scalable AI systems with retrieval, caching, and performance requirements",
            "Automation engineers building autonomous agents for code generation, debugging, and task automation",
        ],
        "curriculum_details": "The course covers practical implementation patterns, performance optimization techniques, and real-world application development using Claude's capabilities within the AWS ecosystem. Topics include API integration, conversation management, prompt engineering, tool use, RAG, advanced features, MCP, and agent architectures.",
    },
    {
        "title": "Claude with Google Cloud's Vertex AI",
        "filename": "11_Claude_with_Google_Vertex_AI.docx",
        "url": "https://anthropic.skilljar.com/claude-with-google-vertex",
        "short_description": "Comprehensive course covering the full spectrum of working with Anthropic models through Google Cloud's Vertex AI.",
        "long_description": "This course provides comprehensive technical training on integrating and deploying Claude AI models through Google Cloud's Vertex AI. Developers learn to implement Claude's API capabilities, from basic request handling to advanced features including tool use, RAG, and MCP. The curriculum covers practical implementation patterns, performance optimization techniques, and production-ready workflows.",
        "course_id": "1k0qj1i9wgzt",
        "published_course_id": "27qd2s5mmsnzm",
        "cost": "Free",
        "executive_summary": "\"Claude with Google Cloud's Vertex AI\" provides comprehensive technical training on deploying Claude through Google Cloud's Vertex AI platform. It mirrors the breadth of the core API course but is specifically tailored for the Google Cloud ecosystem, covering API implementation, conversation management, prompt engineering, tool use, RAG, MCP, and agent-based workflows.",
        "learning_objectives": [
            "Set up and configure Claude models through Google Cloud's Vertex AI",
            "Implement multi-turn conversations with proper message handling and context management",
            "Design and evaluate prompts using systematic testing workflows and automated grading techniques",
            "Apply prompt engineering principles including XML tag structuring, example-based learning, and output control",
            "Build tool-use implementations enabling Claude to interact with external functions and APIs",
            "Develop RAG pipelines using text chunking, embeddings, BM25 search, and contextual retrieval techniques",
            "Utilize advanced Claude features including vision capabilities, PDF processing, citation generation, and prompt caching",
            "Implement the Model Context Protocol for creating custom tools, resources, and prompt templates",
            "Configure and deploy Anthropic Apps including Claude Code for automated development tasks and Computer Use for UI automation",
            "Design agent-based workflows with parallelization, chaining, and routing patterns for complex AI systems",
        ],
        "prerequisites": [
            "Proficiency in Python programming",
            "Experience with Google Cloud Platform",
            "Understanding of JSON data structures",
        ],
        "target_audience": [
            "Backend developers building AI-powered APIs and services",
            "Full-stack engineers integrating LLM capabilities into applications",
            "ML engineers implementing production AI systems",
            "DevOps professionals deploying and scaling Claude implementations",
            "Technical architects designing AI-enhanced system architectures",
            "Developers transitioning from other LLM providers to Claude",
            "Engineers working on document processing, code generation, or automation workflows",
        ],
        "curriculum_details": "The course covers the full spectrum of Claude capabilities on Vertex AI: setup and configuration, conversation management, prompt design and evaluation, prompt engineering, tool use, RAG pipelines, advanced features (vision, PDF, citations, caching), MCP integration, Anthropic Apps deployment, and agent-based workflow design.",
    },
    {
        "title": "Teaching AI Fluency",
        "filename": "12_Teaching_AI_Fluency.docx",
        "url": "https://anthropic.skilljar.com/teaching-ai-fluency",
        "short_description": "Empowers academic faculty, instructional designers, and others to teach and assess AI Fluency in instructor-led settings.",
        "long_description": "This course empowers academic faculty, instructional designers, and others to teach and assess AI Fluency in instructor-led settings. It is a partnership between Anthropic and academic experts Prof. Joseph Feller (University College Cork) and Prof. Rick Dakan (Ringling College). The course builds upon the foundational 4D AI Fluency Framework.",
        "course_id": "zjx5xtwwhlir",
        "published_course_id": "1z1ggntgogc6y",
        "cost": "Free",
        "executive_summary": "\"Teaching AI Fluency\" is designed for educators who want to teach and assess AI fluency in instructor-led classroom settings. Developed with academic experts, it provides pedagogical strategies for delivering AI fluency education based on the 4D Framework. This course complements the self-paced AI Fluency courses with instructor-facilitated approaches.",
        "learning_objectives": [],
        "prerequisites": [
            "Understanding of the 4D AI Fluency Framework from the core foundational course",
        ],
        "target_audience": [
            "Academic faculty",
            "Instructional designers",
            "Educators teaching AI fluency in instructor-led settings",
        ],
        "curriculum_details": "The course provides strategies and resources for teaching and assessing AI Fluency in instructor-led environments. It builds on the 4D AI Fluency Framework and includes supplementary materials.",
        "resources": [
            "About this course (PDF)",
            "Meet your instructors (PDF)",
            "How we used AI in building the course (PDF)",
        ],
        "certification": "After finishing the course, participants have the opportunity to take a final assessment and receive a certificate of completion.",
    },
    {
        "title": "AI Fluency for Nonprofits",
        "filename": "13_AI_Fluency_for_Nonprofits.docx",
        "url": "https://anthropic.skilljar.com/ai-fluency-for-nonprofits",
        "short_description": "Empowers nonprofit professionals to develop AI fluency to increase organizational impact and efficiency while staying true to their mission and values.",
        "long_description": "This course empowers nonprofit professionals to develop AI fluency in order to increase organizational impact and efficiency while staying true to their mission and values. It teaches practical collaboration skills via the 4D Framework covering Delegation, Description, Discernment, and Diligence, applied to common nonprofit functions. This is a partnership between Anthropic and GivingTuesday, grounded in research about nonprofit professional needs.",
        "course_id": "37f6rpi1f0h0",
        "published_course_id": "2qv34d09tdzk",
        "cost": "Free",
        "executive_summary": "\"AI Fluency for Nonprofits\" adapts the 4D AI Fluency Framework specifically for nonprofit organizations. Developed in partnership with GivingTuesday, it addresses the unique considerations of limited resources, multiple stakeholder accountabilities, and mission-driven priorities. The course covers practical AI applications across fundraising, communications, program delivery, operations, and leadership.",
        "learning_objectives": [],
        "prerequisites": [
            "Recommended: Complete \"AI Fluency: Framework & Foundations\" course first for deeper conceptual understanding",
            "Access to an AI chat tool for hands-on practice (Claude.ai recommended, though any chatbot works)",
        ],
        "target_audience": [
            "Nonprofit professionals in fundraising",
            "Nonprofit professionals in communications",
            "Nonprofit professionals in program delivery",
            "Nonprofit professionals in operations",
            "Nonprofit leaders and managers",
        ],
        "curriculum_details": "The course applies the 4D Framework (Delegation, Description, Discernment, Diligence) to nonprofit-specific scenarios. It explores how AI can support fundraising, communications, program delivery, operations, and leadership while maintaining mission alignment and values. The curriculum is grounded in research about nonprofit professional needs, concerns, and AI adoption aspirations.",
        "topics": [
            "4D Framework applied to nonprofit contexts: Delegation, Description, Discernment, and Diligence",
            "AI applications in fundraising",
            "AI applications in communications",
            "AI applications in program delivery",
            "AI applications in operations",
            "AI applications in nonprofit leadership",
            "Responsible AI use aligned with mission and values",
        ],
    },
    {
        "title": "Introduction to Agent Skills",
        "filename": "14_Introduction_to_Agent_Skills.docx",
        "url": "https://anthropic.skilljar.com/introduction-to-agent-skills",
        "short_description": "Learn how to build, configure, and share Skills in Claude Code -- reusable markdown instructions that Claude automatically applies to the right tasks at the right time.",
        "long_description": "This comprehensive course covers building reusable instruction sets (Skills) for Claude Code. Learners progress from foundational concepts through distribution strategies, covering skill creation, trigger descriptions, directory organization, advanced configuration, team distribution, enterprise deployment, and subagent integration.",
        "course_id": "377fuwy2g2o8k",
        "published_course_id": "lczzcndyyoct",
        "cost": "Free",
        "executive_summary": "\"Introduction to Agent Skills\" teaches developers how to build, configure, and distribute reusable markdown-based Skills for Claude Code. Skills are automatically applied to the right tasks at the right time, enabling consistent workflows and team standardization. The course covers the full lifecycle from creation to enterprise deployment.",
        "learning_objectives": [
            "Understanding Skills versus alternative customization methods (CLAUDE.md, hooks, subagents)",
            "Creating Skills with proper SKILL.md frontmatter formatting",
            "Writing effective trigger descriptions for reliable skill matching",
            "Organizing skill directories with progressive disclosure for context efficiency",
            "Advanced configuration including tool access restrictions and script execution",
            "Team-based distribution through repository commits",
            "Broader organization distribution via plugins",
            "Enterprise deployment using managed settings",
            "Integrating Skills into custom subagents for specialized task delegation",
            "Comprehensive troubleshooting guidance addressing trigger failures, priority conflicts, and runtime errors",
        ],
        "prerequisites": [],
        "target_audience": [
            "Developers building consistent Claude Code workflows",
            "Teams seeking to standardize AI-assisted development practices",
            "Enterprise organizations deploying Claude Code at scale",
        ],
        "curriculum_details": "The course employs progressive skill building, moving from individual skill creation to team workflows, organizational distribution, and advanced integration patterns. It covers SKILL.md formatting, trigger descriptions, directory organization, advanced configuration, distribution strategies (repository, plugins, managed settings), subagent integration, and troubleshooting.",
    },
    {
        "title": "Introduction to Subagents",
        "filename": "15_Introduction_to_Subagents.docx",
        "url": "https://anthropic.skilljar.com/introduction-to-subagents",
        "short_description": "Use and create sub-agents in Claude Code to manage context, delegate tasks, and build specialized workflows.",
        "long_description": "This course teaches participants how to use and create sub-agents in Claude Code to manage context, delegate tasks, and build specialized workflows that keep the main conversation clean and focused.",
        "course_id": "305ppx2mtvlia",
        "published_course_id": "3h0rqpiust8ut",
        "cost": "Free",
        "executive_summary": "\"Introduction to Subagents\" teaches developers how to leverage sub-agents in Claude Code for context management, task delegation, and specialized workflow construction. It covers sub-agent architecture, custom creation, design patterns, and decision-making frameworks for when delegation is appropriate.",
        "learning_objectives": [
            {"title": "How Sub-Agents Work", "description": "Understanding what happens when Claude Code spins up a separate context window, how inputs flow in, and how summaries come back"},
            {"title": "Creating Custom Sub-Agents", "description": "Using the /agents command to build sub-agents tailored to your workflow, from code reviewers to documentation generators"},
            {"title": "Designing Effective Sub-Agents", "description": "Patterns that make sub-agents reliable, including structured output formats, obstacle reporting, and limiting tool access"},
            {"title": "When to Use Them (and When Not To)", "description": "Practical guidance on where sub-agents help the most and the common anti-patterns to avoid"},
        ],
        "prerequisites": [],
        "target_audience": [
            "Claude Code users seeking to manage complex, multi-step workflows",
            "Developers building specialized AI-assisted development pipelines",
        ],
        "curriculum_details": "The course covers sub-agent architecture and context flow, custom sub-agent creation via the /agents command, design patterns for reliability (structured outputs, obstacle reporting, tool access limits), and decision frameworks for appropriate sub-agent use. By the end, learners know how to break complex work into focused pieces and build sub-agents that finish on time and report back clearly.",
    },
    {
        "title": "AI Capabilities and Limitations",
        "filename": "16_AI_Capabilities_and_Limitations.docx",
        "url": "https://anthropic.skilljar.com/ai-capabilities-and-limitations",
        "short_description": "A working mental model of how modern generative AI systems behave and why.",
        "long_description": "This introductory course gives learners a working mental model of how modern generative AI systems behave and why. It's the companion to the 4D Framework: where that course teaches the human competencies (Delegation, Description, Discernment, Diligence), this one teaches the machine properties those competencies are responding to.",
        "course_id": "3oa2jr8dqiak3",
        "published_course_id": "3g5wgfrtrwkta",
        "cost": "Free",
        "executive_summary": "\"AI Capabilities and Limitations\" provides a foundational mental model of how modern generative AI systems work and why they behave as they do. It serves as the technical companion to the 4D AI Fluency Framework, teaching the machine properties that the human competencies (Delegation, Description, Discernment, Diligence) respond to. This course bridges the gap between practical AI use and understanding AI behavior.",
        "learning_objectives": [
            "Recognize which kind of unexpected an AI output is when encountering unexpected results",
            "Locate roughly where on a capability-to-limitation continuum a given task lands",
            "Respond with targeted fixes to unexpected AI outputs",
        ],
        "prerequisites": [],
        "target_audience": [
            "Learners seeking to understand AI system behavior",
            "Anyone looking for a technical foundation to complement the 4D AI Fluency Framework",
        ],
        "curriculum_details": "The course teaches a working mental model of generative AI behavior, focusing on the capability-to-limitation continuum. It covers recognizing types of unexpected outputs, diagnosing where tasks fall on the capability spectrum, and applying targeted fixes. It complements the 4D Framework by addressing the machine side of human-AI collaboration.",
    },
]

# Generate all reports
print("=" * 60)
print("  ANTHROPIC ACADEMY - Professional Course Reports")
print("  Generating 16 .docx files with enhanced formatting...")
print("=" * 60)
for course in courses:
    create_report(course)
print("=" * 60)
print(f"  All {len(courses)} reports generated successfully!")
print(f"  Output: {OUTPUT_DIR}")
print("=" * 60)
